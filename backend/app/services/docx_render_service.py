"""
docxtpl 渲染服务
使用 docxtpl 库渲染 Word 模板，生成最终文档
"""
import io
import re
from typing import Dict, Any, Optional
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn


class DocxRenderService:
    """docxtpl 渲染服务"""
    
    def __init__(self):
        pass
    
    def _fix_placeholder_format(self, template_bytes: bytes) -> bytes:
        """
        修复模板中的占位符格式错误
        
        Args:
            template_bytes: 原始模板文件
            
        Returns:
            修复后的模板文件
        """
        try:
            doc = Document(io.BytesIO(template_bytes))
            
            fixed_count = 0
            
            def fix_text(text: str) -> str:
                """修复文本中的占位符格式"""
                if not text:
                    return text
                
                original = text
                
                # 1. 先修复三个大括号 -> 两个大括号
                fixed = text.replace("{{{", "{{").replace("}}}", "}}")
                
                # 2. 修复不完整的占位符（但不要修改已经正确的）
                # 只修复明显错误的情况
                
                # 情况 A: {{ variable } (缺少一个结束括号) - 但不是 {{ variable }}
                # 使用负向后顾确保后面不是 }}
                fixed = re.sub(r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}(?!\})', r'{{ \1 }}', fixed)
                
                # 情况 B: { variable }} (缺少一个开始括号) - 但不是 {{ variable }}
                # 使用负向前瞻确保前面不是 {{
                fixed = re.sub(r'(?<!\{)\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}', r'{{ \1 }}', fixed)
                
                # 3. 不要再添加其他修复逻辑，避免误伤正确的占位符
                
                if fixed != original:
                    return fixed
                
                return text
            
            def fix_paragraph(para):
                """修复段落中的占位符 - 在段落级别处理"""
                nonlocal fixed_count
                
                # 获取完整的段落文本
                full_text = para.text
                
                # 如果没有占位符标记，跳过
                if "{" not in full_text and "}" not in full_text:
                    return
                
                # 修复完整文本
                fixed_text = fix_text(full_text)
                
                if fixed_text != full_text:
                    print(f"[DocxRender] Fixing paragraph:")
                    print(f"  Before: {full_text[:100]}")
                    print(f"  After:  {fixed_text[:100]}")
                    
                    # 保存第一个 run 的格式
                    if para.runs:
                        first_run = para.runs[0]
                        font_info = {
                            'name': first_run.font.name,
                            'size': first_run.font.size,
                            'bold': first_run.font.bold,
                            'italic': first_run.font.italic,
                            'underline': first_run.font.underline,
                        }
                        
                        # 清空所有 runs
                        for run in para.runs:
                            run.text = ""
                        
                        # 将修复后的文本放在第一个 run 中
                        first_run.text = fixed_text
                        
                        # 恢复格式
                        if font_info['name']:
                            first_run.font.name = font_info['name']
                        if font_info['size']:
                            first_run.font.size = font_info['size']
                        if font_info['bold'] is not None:
                            first_run.font.bold = font_info['bold']
                        if font_info['italic'] is not None:
                            first_run.font.italic = font_info['italic']
                        if font_info['underline'] is not None:
                            first_run.font.underline = font_info['underline']
                        
                        fixed_count += 1
            
            # 修复段落
            for para in doc.paragraphs:
                fix_paragraph(para)
            
            # 修复表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            fix_paragraph(para)
            
            # 修复页眉页脚
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        fix_paragraph(para)
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        fix_paragraph(para)
            
            if fixed_count > 0:
                print(f"[DocxRender] Total paragraphs fixed: {fixed_count}")
                # 保存修复后的文档
                buffer = io.BytesIO()
                doc.save(buffer)
                return buffer.getvalue()
            else:
                print(f"[DocxRender] No placeholder fixes needed")
                return template_bytes
                
        except Exception as e:
            print(f"[DocxRender] Error fixing placeholders: {e}")
            import traceback
            traceback.print_exc()
            # 如果修复失败，返回原始文件
            return template_bytes
    
    async def render_template(
        self,
        template_bytes: bytes,
        context: Dict[str, Any]
    ) -> bytes:
        """
        渲染 Word 模板
        
        Args:
            template_bytes: 模板文件二进制内容
            context: 渲染数据字典
            
        Returns:
            渲染后的 Word 文档二进制
        """
        try:
            # 暂时禁用修复逻辑，直接使用原始模板
            # print(f"[DocxRender] Fixing placeholder format...")
            # fixed_template_bytes = self._fix_placeholder_format(template_bytes)
            
            # 加载模板
            doc = DocxTemplate(io.BytesIO(template_bytes))
            
            # 调试：打印模板中的占位符
            try:
                template_vars = doc.get_undeclared_template_variables()
                print(f"[DocxRender] Template variables: {template_vars}")
            except Exception as e:
                print(f"[DocxRender] Warning: Could not get template variables: {e}")
            
            print(f"[DocxRender] Context keys: {list(context.keys())}")
            print(f"[DocxRender] Context values preview: {str(context)[:500]}")
            
            # 渲染
            doc.render(context)
            
            # 保存到内存
            buffer = io.BytesIO()
            doc.save(buffer)
            
            print(f"[DocxRender] Render successful: {len(buffer.getvalue())} bytes")
            
            return buffer.getvalue()
            
        except Exception as e:
            print(f"[DocxRender] Render error: {e}")
            print(f"[DocxRender] Error type: {type(e).__name__}")
            import traceback
            traceback.print_exc()
            raise Exception(f"模板渲染失败: {str(e)}")
    
    async def validate_context(
        self,
        fields: Dict[str, Any],
        context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        验证渲染数据是否完整
        
        Args:
            fields: 模板字段定义
            context: 待验证的数据字典
            
        Returns:
            验证结果
        """
        missing_fields = []
        
        for field_name, field_info in fields.items():
            if field_info.get("required", False):
                if field_name not in context or not context[field_name]:
                    missing_fields.append({
                        "name": field_name,
                        "label": field_info.get("label", field_name)
                    })
        
        return {
            "valid": len(missing_fields) == 0,
            "missing_fields": missing_fields
        }
    
    def get_template_variables(self, template_bytes: bytes) -> list:
        """
        获取模板中的所有变量名
        
        Args:
            template_bytes: 模板文件二进制
            
        Returns:
            变量名列表
        """
        try:
            doc = DocxTemplate(io.BytesIO(template_bytes))
            # 使用 docxtpl 的内置方法获取未声明的变量
            variables = doc.get_undeclared_template_variables()
            return list(variables)
        except Exception as e:
            print(f"[DocxRender] Get variables error: {e}")
            return []


# 创建全局实例
docx_render_service = DocxRenderService()
