"""
文档导出服务
支持将文书导出为 PDF 和 DOCX 格式
"""
import io
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import traceback


class DocumentExportService:
    """文档导出服务类"""
    
    def __init__(self):
        # 注册中文字体（如果有的话）
        try:
            # 尝试注册常见的中文字体
            # 实际部署时需要确保字体文件存在
            pass
        except:
            pass
    
    async def export_to_pdf(
        self,
        content: str,
        title: str,
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        导出为 PDF 格式
        
        Args:
            content: 文档内容
            title: 文档标题
            options: 导出选项（水印、密级等）
            
        Returns:
            PDF 文件的二进制内容
        """
        try:
            options = options or {}
            
            # 创建 PDF 文档
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # 准备样式
            styles = getSampleStyleSheet()
            
            # 标题样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=RGBColor(0, 0, 0),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # 正文样式
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=12,
                leading=20,
                alignment=TA_JUSTIFY,
                fontName='Helvetica'
            )
            
            # 构建文档内容
            story = []
            
            # 添加标题
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # 添加密级标注（如果有）
            if options.get('security_level'):
                security_style = ParagraphStyle(
                    'Security',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=RGBColor(255, 0, 0),
                    alignment=TA_CENTER
                )
                story.append(Paragraph(f"密级：{options['security_level']}", security_style))
                story.append(Spacer(1, 0.1 * inch))
            
            # 添加正文内容
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    # 简单的 HTML 转义
                    para_text = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(para_text, body_style))
                    story.append(Spacer(1, 0.1 * inch))
            
            # 添加水印（如果需要）
            if options.get('watermark'):
                # 简单实现：在末尾添加水印文本
                watermark_style = ParagraphStyle(
                    'Watermark',
                    parent=styles['Normal'],
                    fontSize=8,
                    textColor=RGBColor(200, 200, 200),
                    alignment=TA_CENTER
                )
                story.append(Spacer(1, 0.5 * inch))
                story.append(Paragraph(options['watermark'], watermark_style))
            
            # 生成 PDF
            doc.build(story)
            
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            print(f"[Export] PDF generated successfully: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except Exception as e:
            print(f"[Export] PDF generation error: {e}")
            traceback.print_exc()
            raise
    
    async def export_to_docx(
        self,
        content: str,
        title: str,
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        导出为 DOCX 格式
        
        Args:
            content: 文档内容
            title: 文档标题
            options: 导出选项（水印、密级等）
            
        Returns:
            DOCX 文件的二进制内容
        """
        try:
            options = options or {}
            
            # 创建 Word 文档
            document = Document()
            
            # 设置文档属性
            document.core_properties.title = title
            
            # 添加标题
            title_para = document.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加密级标注（如果有）
            if options.get('security_level'):
                security_para = document.add_paragraph()
                security_run = security_para.add_run(f"密级：{options['security_level']}")
                security_run.font.size = Pt(10)
                security_run.font.color.rgb = RGBColor(255, 0, 0)
                security_run.bold = True
                security_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph()  # 空行
            
            # 添加正文内容
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = document.add_paragraph(para_text)
                    # 设置段落格式
                    para_format = para.paragraph_format
                    para_format.line_spacing = 1.5
                    para_format.space_after = Pt(6)
                    
                    # 设置字体
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Arial'
            
            # 添加水印（如果需要）
            if options.get('watermark'):
                document.add_paragraph()  # 空行
                watermark_para = document.add_paragraph(options['watermark'])
                watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in watermark_para.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(200, 200, 200)
            
            # 保存到内存
            buffer = io.BytesIO()
            document.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            print(f"[Export] DOCX generated successfully: {len(docx_bytes)} bytes")
            return docx_bytes
            
        except Exception as e:
            print(f"[Export] DOCX generation error: {e}")
            traceback.print_exc()
            raise
    
    async def export_document(
        self,
        content: str,
        title: str,
        format: str = 'pdf',
        options: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        导出文档（统一接口）
        
        Args:
            content: 文档内容
            title: 文档标题
            format: 导出格式（pdf 或 docx）
            options: 导出选项
            
        Returns:
            文档的二进制内容
        """
        if format.lower() == 'pdf':
            return await self.export_to_pdf(content, title, options)
        elif format.lower() in ['docx', 'doc']:
            return await self.export_to_docx(content, title, options)
        else:
            raise ValueError(f"Unsupported export format: {format}")


# 创建全局实例
document_export_service = DocumentExportService()
