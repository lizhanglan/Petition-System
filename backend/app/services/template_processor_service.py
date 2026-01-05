"""
Word 模板处理服务
负责解析用户上传的 Word 文档，调用 AI 识别可变字段，并自动替换为占位符
"""
import io
import re
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from app.services.deepseek_service import deepseek_service


class TemplateProcessorService:
    """模板处理服务"""
    
    def __init__(self):
        pass
    
    async def parse_document(self, file_bytes: bytes) -> Dict[str, Any]:
        """
        解析 Word 文档内容
        
        Args:
            file_bytes: Word 文档的二进制内容
            
        Returns:
            包含文档文本和结构信息的字典
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # 提取所有文本
            full_text = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return {
                "success": True,
                "text": "\n".join(full_text),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    async def identify_fields(self, document_text: str) -> Dict[str, Any]:
        """
        调用 AI 识别文档中的可变字段
        
        Args:
            document_text: 文档纯文本内容
            
        Returns:
            字段识别结果
        """
        result = await deepseek_service.identify_template_fields(document_text)
        return result
    
    async def replace_with_placeholders(
        self, 
        file_bytes: bytes, 
        replacements: Dict[str, str]
    ) -> Tuple[bytes, bool]:
        """
        在 Word 文档中替换文本为占位符
        
        Args:
            file_bytes: 原始 Word 文档
            replacements: 替换映射 {原文本: 占位符}
            
        Returns:
            (处理后的文档二进制, 是否成功)
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # 按原文本长度降序排列，先替换长的字符串，防止短串误匹配
            sorted_replacements = dict(
                sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)
            )
            
            print(f"[TemplateProcessor] Replacement order: {list(sorted_replacements.keys())}")
            
            # 替换段落中的文本
            for para in doc.paragraphs:
                self._replace_in_paragraph(para, sorted_replacements)
            
            # 替换表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, sorted_replacements)
            
            # 替换页眉页脚
            for section in doc.sections:
                # 页眉
                if section.header:
                    for para in section.header.paragraphs:
                        self._replace_in_paragraph(para, sorted_replacements)
                # 页脚
                if section.footer:
                    for para in section.footer.paragraphs:
                        self._replace_in_paragraph(para, sorted_replacements)
            
            # 保存到内存
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue(), True
            
        except Exception as e:
            print(f"[TemplateProcessor] Replace error: {e}")
            return file_bytes, False
    
    def _replace_in_paragraph(self, paragraph: Paragraph, replacements: Dict[str, str]):
        """
        在段落中替换文本，尽可能保留格式
        
        注意：replacements 应该已按长度降序排列
        """
        for old_text, new_text in replacements.items():
            if old_text not in paragraph.text:
                continue
            
            print(f"[TemplateProcessor] Replacing '{old_text}' with '{new_text}' in paragraph")
            
            # 验证新文本格式（确保是正确的 Jinja2 格式）
            if not new_text.startswith("{{") or not new_text.endswith("}}"):
                print(f"[TemplateProcessor] WARNING: Invalid placeholder format: {new_text}")
                # 修正格式
                if new_text.startswith("{") and not new_text.startswith("{{"):
                    new_text = "{" + new_text
                if new_text.endswith("}") and not new_text.endswith("}}"):
                    new_text = new_text + "}"
                print(f"[TemplateProcessor] Corrected to: {new_text}")
            
            # 尝试在单个 run 中替换（最佳情况）
            replaced_in_run = False
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replaced_in_run = True
                    print(f"[TemplateProcessor] Replaced in single run successfully")
            
            # 如果在单个 run 中成功替换，继续下一个
            if replaced_in_run:
                continue
            
            # 跨 run 的情况：需要特殊处理以保留格式
            if old_text in paragraph.text:
                print(f"[TemplateProcessor] Cross-run replacement needed")
                
                # 保存所有 runs 的格式信息
                runs_info = []
                for run in paragraph.runs:
                    runs_info.append({
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_bold': run.font.bold,
                        'font_italic': run.font.italic,
                        'font_underline': run.font.underline,
                        'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                    })
                
                # 获取完整文本并执行替换
                full_text = paragraph.text
                new_full_text = full_text.replace(old_text, new_text)
                
                # 验证替换后的文本中没有错误的占位符格式
                if "{{{" in new_full_text or "}}}" in new_full_text:
                    print(f"[TemplateProcessor] ERROR: Triple braces detected in: {new_full_text[:100]}")
                    new_full_text = new_full_text.replace("{{{", "{{").replace("}}}", "}}")
                    print(f"[TemplateProcessor] Corrected triple braces")
                
                # 清空所有 runs
                for run in paragraph.runs:
                    run.text = ""
                
                # 策略：将新文本完整地放在第一个 run 中，保留其格式
                if paragraph.runs and runs_info:
                    first_run = paragraph.runs[0]
                    first_run.text = new_full_text
                    
                    # 恢复第一个 run 的格式
                    first_format = runs_info[0]
                    if first_format['font_name']:
                        first_run.font.name = first_format['font_name']
                    if first_format['font_size']:
                        first_run.font.size = first_format['font_size']
                    if first_format['font_bold'] is not None:
                        first_run.font.bold = first_format['font_bold']
                    if first_format['font_italic'] is not None:
                        first_run.font.italic = first_format['font_italic']
                    if first_format['font_underline'] is not None:
                        first_run.font.underline = first_format['font_underline']
                    if first_format['font_color']:
                        first_run.font.color.rgb = first_format['font_color']
                
                print(f"[TemplateProcessor] Cross-run replacement completed, text now in single run")
    
    async def process_template(
        self, 
        file_bytes: bytes, 
        file_name: str
    ) -> Dict[str, Any]:
        """
        完整的模板处理流程
        
        Args:
            file_bytes: 上传的 Word 文档
            file_name: 文件名
            
        Returns:
            处理结果，包含识别的字段和处理后的模板
        """
        # 1. 解析文档
        parse_result = await self.parse_document(file_bytes)
        if not parse_result.get("success"):
            return {
                "success": False,
                "error": f"文档解析失败: {parse_result.get('error')}"
            }
        
        document_text = parse_result["text"]
        
        # 2. AI 识别字段
        ai_result = await self.identify_fields(document_text)
        if not ai_result or not ai_result.get("success"):
            return {
                "success": False,
                "error": "AI 识别字段失败"
            }
        
        fields = ai_result.get("fields", {})
        replacements = ai_result.get("replacements", {})
        
        if not replacements:
            return {
                "success": False,
                "error": "未识别到可变字段"
            }
        
        # 3. 替换占位符
        template_bytes, replace_success = await self.replace_with_placeholders(
            file_bytes, replacements
        )
        
        if not replace_success:
            return {
                "success": False,
                "error": "占位符替换失败"
            }
        
        return {
            "success": True,
            "fields": fields,
            "replacements": replacements,
            "template_bytes": template_bytes,
            "original_text": document_text
        }


# 创建全局实例
template_processor_service = TemplateProcessorService()
