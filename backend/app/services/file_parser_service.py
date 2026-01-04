"""
文件解析服务
支持 PDF 和 Word 文档的文本提取和格式保留
"""
import io
from typing import Optional, Dict, Any
from docx import Document
from PyPDF2 import PdfReader
import traceback


class FileParserService:
    """文件解析服务类"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'doc', 'docx']
    
    async def parse_file(self, file_content: bytes, file_type: str) -> Optional[Dict[str, Any]]:
        """
        解析文件内容
        
        Args:
            file_content: 文件二进制内容
            file_type: 文件类型（pdf, doc, docx）
            
        Returns:
            解析结果字典，包含：
            - text: 提取的文本内容
            - metadata: 文件元数据
            - structure: 文档结构信息
        """
        try:
            if file_type.lower() == 'pdf':
                return await self._parse_pdf(file_content)
            elif file_type.lower() in ['doc', 'docx']:
                return await self._parse_word(file_content)
            else:
                print(f"[FileParser] Unsupported file type: {file_type}")
                return None
        except Exception as e:
            print(f"[FileParser] Error parsing file: {e}")
            traceback.print_exc()
            return None
    
    async def _parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        解析 PDF 文件
        
        Args:
            file_content: PDF 文件二进制内容
            
        Returns:
            解析结果字典
        """
        try:
            # 创建 PDF 读取器
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            # 提取元数据
            metadata = {
                'num_pages': len(pdf_reader.pages),
                'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else '',
            }
            
            # 提取文本内容
            text_content = []
            structure = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        structure.append({
                            'page': page_num,
                            'text_length': len(page_text),
                            'has_content': bool(page_text.strip())
                        })
                except Exception as e:
                    print(f"[FileParser] Error extracting page {page_num}: {e}")
                    structure.append({
                        'page': page_num,
                        'error': str(e),
                        'has_content': False
                    })
            
            full_text = '\n\n'.join(text_content)
            
            print(f"[FileParser] PDF parsed successfully: {len(pdf_reader.pages)} pages, {len(full_text)} characters")
            
            return {
                'text': full_text,
                'metadata': metadata,
                'structure': structure,
                'format': 'pdf'
            }
            
        except Exception as e:
            print(f"[FileParser] PDF parsing error: {e}")
            traceback.print_exc()
            raise
    
    async def _parse_word(self, file_content: bytes) -> Dict[str, Any]:
        """
        解析 Word 文档
        
        Args:
            file_content: Word 文件二进制内容
            
        Returns:
            解析结果字典
        """
        try:
            # 检查是否为旧版 .doc 格式
            doc_file = io.BytesIO(file_content)
            
            # 尝试检测文件头
            file_header = file_content[:8]
            if file_header[:4] == b'\xd0\xcf\x11\xe0':  # OLE2 格式（旧版 .doc）
                print(f"[FileParser] 检测到旧版 Word 格式（.doc），不支持解析")
                raise ValueError("不支持旧版 Word 格式（.doc），请转换为 .docx 格式后重试")
            
            # 创建 Word 文档对象
            document = Document(doc_file)
            
            # 提取元数据
            core_properties = document.core_properties
            metadata = {
                'author': core_properties.author or '',
                'title': core_properties.title or '',
                'subject': core_properties.subject or '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
            }
            
            # 提取文本内容和结构
            text_content = []
            structure = []
            
            # 提取段落
            for para_num, paragraph in enumerate(document.paragraphs, 1):
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    structure.append({
                        'type': 'paragraph',
                        'index': para_num,
                        'text_length': len(paragraph.text),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # 提取表格
            for table_num, table in enumerate(document.tables, 1):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    table_content = '\n'.join(table_text)
                    text_content.append(f"\n[表格 {table_num}]\n{table_content}\n")
                    structure.append({
                        'type': 'table',
                        'index': table_num,
                        'rows': len(table.rows),
                        'columns': len(table.columns) if table.rows else 0
                    })
            
            full_text = '\n'.join(text_content)
            
            print(f"[FileParser] Word document parsed successfully: {len(document.paragraphs)} paragraphs, {len(document.tables)} tables, {len(full_text)} characters")
            
            return {
                'text': full_text,
                'metadata': metadata,
                'structure': structure,
                'format': 'docx'
            }
            
        except Exception as e:
            print(f"[FileParser] Word parsing error: {e}")
            traceback.print_exc()
            raise
    
    def extract_key_info(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        从解析结果中提取关键信息
        
        Args:
            parsed_data: 解析结果
            
        Returns:
            关键信息字典
        """
        text = parsed_data.get('text', '')
        
        # 基础统计
        key_info = {
            'total_length': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.split('\n')),
            'has_tables': any(item.get('type') == 'table' for item in parsed_data.get('structure', [])),
            'format': parsed_data.get('format', 'unknown')
        }
        
        # 尝试提取常见的文书要素（简单规则）
        lines = text.split('\n')
        for line in lines[:10]:  # 检查前10行
            line = line.strip()
            if '文号' in line or '编号' in line:
                key_info['document_number'] = line
            elif '主送' in line:
                key_info['recipient'] = line
            elif '日期' in line or '年' in line and '月' in line and '日' in line:
                key_info['date'] = line
        
        return key_info


# 创建全局实例
file_parser_service = FileParserService()
